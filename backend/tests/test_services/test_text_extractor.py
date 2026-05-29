import os
import tempfile

import pytest
from docx import Document as DocxDocument

from app.services.text_extractor import extract_text


@pytest.fixture()
def txt_file():
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write("This is a test text file.\nIt has two lines.")
        name = f.name
    yield name
    os.unlink(name)


@pytest.fixture()
def docx_file():
    doc = DocxDocument()
    doc.add_paragraph("This is a test DOCX file.")
    doc.add_paragraph("It has two paragraphs.")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        name = f.name
    doc.save(name)
    yield name
    os.unlink(name)


def test_extract_txt(txt_file):
    text = extract_text(txt_file, "text/plain")
    assert "test text file" in text
    assert "two lines" in text


def test_extract_txt_by_extension(txt_file):
    text = extract_text(txt_file, "application/octet-stream")
    assert "test text file" in text


def test_extract_docx(docx_file):
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    text = extract_text(docx_file, mime)
    assert "test DOCX file" in text
    assert "two paragraphs" in text


def test_extract_docx_by_extension(docx_file):
    text = extract_text(docx_file, "application/octet-stream")
    assert "test DOCX file" in text
